"""
Document Processor
Handles extraction of text from PDF, TXT, CSV, and DOCX files.
Supports OCR for scanned PDFs using pytesseract.
"""

import io
import os
import logging
from pathlib import Path
from typing import Optional

import pandas as pd

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Unified document processing engine.
    Extracts clean text content from multiple file formats.
    """

    SUPPORTED_FORMATS = {
        "pdf": "application/pdf",
        "txt": "text/plain",
        "csv": "text/csv",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    def __init__(self, ocr_enabled: bool = False):
        self.ocr_enabled = ocr_enabled
        self._check_dependencies()

    def _check_dependencies(self):
        """Check optional dependencies availability."""
        self._has_pymupdf = False
        self._has_docx = False
        self._has_tesseract = False

        try:
            import fitz  # PyMuPDF
            self._has_pymupdf = True
        except ImportError:
            logger.warning("PyMuPDF not available. PDF support limited.")

        try:
            import docx
            self._has_docx = True
        except ImportError:
            logger.warning("python-docx not available. DOCX support disabled.")

        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            self._has_tesseract = True
        except Exception:
            logger.warning("Tesseract not available. OCR support disabled.")

    def process(self, file_obj, filename: str) -> dict:
        """
        Main processing method. Routes to appropriate handler based on file type.

        Returns:
            dict: {
                'text': str,
                'metadata': dict,
                'pages': int,
                'word_count': int,
                'format': str,
                'success': bool,
                'error': Optional[str]
            }
        """
        ext = Path(filename).suffix.lower().strip(".")

        processors = {
            "pdf": self._process_pdf,
            "txt": self._process_txt,
            "csv": self._process_csv,
            "docx": self._process_docx,
        }

        if ext not in processors:
            return {
                "text": "",
                "metadata": {},
                "pages": 0,
                "word_count": 0,
                "format": ext,
                "success": False,
                "error": f"Unsupported file format: .{ext}",
            }

        try:
            result = processors[ext](file_obj, filename)
            result["format"] = ext
            result["success"] = True
            result["error"] = None
            result["word_count"] = len(result.get("text", "").split())
            return result
        except Exception as e:
            logger.error(f"Error processing {filename}: {e}")
            return {
                "text": "",
                "metadata": {"filename": filename},
                "pages": 0,
                "word_count": 0,
                "format": ext,
                "success": False,
                "error": str(e),
            }

    def _process_pdf(self, file_obj, filename: str) -> dict:
        """Extract text from PDF using PyMuPDF with OCR fallback."""
        if not self._has_pymupdf:
            raise ImportError("PyMuPDF is required for PDF processing. Run: pip install PyMuPDF")

        import fitz

        if hasattr(file_obj, "read"):
            pdf_bytes = file_obj.read()
        else:
            pdf_bytes = file_obj

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        full_text = []
        page_count = len(doc)
        ocr_pages = 0

        for page_num, page in enumerate(doc):
            text = page.get_text("text")

            # If page has very little text, try OCR
            if len(text.strip()) < 50 and self.ocr_enabled and self._has_tesseract:
                try:
                    import pytesseract
                    from PIL import Image

                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    text = pytesseract.image_to_string(img)
                    ocr_pages += 1
                except Exception as ocr_err:
                    logger.warning(f"OCR failed on page {page_num + 1}: {ocr_err}")

            full_text.append(f"--- Page {page_num + 1} ---\n{text}")

        doc.close()

        return {
            "text": "\n".join(full_text),
            "metadata": {
                "filename": filename,
                "pages": page_count,
                "ocr_pages": ocr_pages,
                "ocr_enabled": self.ocr_enabled,
            },
            "pages": page_count,
        }

    def _process_txt(self, file_obj, filename: str) -> dict:
        """Extract text from plain text files with encoding detection."""
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252", "ascii"]

        if hasattr(file_obj, "read"):
            raw_bytes = file_obj.read()
        else:
            raw_bytes = file_obj

        text = None
        for encoding in encodings:
            try:
                text = raw_bytes.decode(encoding)
                break
            except (UnicodeDecodeError, AttributeError):
                continue

        if text is None:
            text = raw_bytes.decode("utf-8", errors="replace")

        return {
            "text": text,
            "metadata": {"filename": filename, "encoding": encoding},
            "pages": 1,
        }

    def _process_csv(self, file_obj, filename: str) -> dict:
        """Process CSV files — extract all text content including headers and values."""
        try:
            df = pd.read_csv(file_obj)
        except Exception:
            # Try different separators
            file_obj.seek(0)
            for sep in [";", "\t", "|"]:
                try:
                    file_obj.seek(0)
                    df = pd.read_csv(file_obj, sep=sep)
                    break
                except Exception:
                    continue
            else:
                file_obj.seek(0)
                content = file_obj.read()
                if isinstance(content, bytes):
                    content = content.decode("utf-8", errors="replace")
                return {
                    "text": content,
                    "metadata": {"filename": filename},
                    "pages": 1,
                }

        # Build readable text representation of CSV
        text_parts = []
        text_parts.append(f"CSV File: {filename}")
        text_parts.append(f"Columns ({len(df.columns)}): {', '.join(df.columns.tolist())}")
        text_parts.append(f"Rows: {len(df)}")
        text_parts.append("\n--- Data Contents ---")

        # Full CSV as text for detection
        text_parts.append(df.to_string(index=False))

        # Also add raw values concatenated for better pattern matching
        text_parts.append("\n--- Raw Values ---")
        for col in df.columns:
            text_parts.append(f"\n[Column: {col}]")
            text_parts.append(df[col].astype(str).str.cat(sep="\n"))

        return {
            "text": "\n".join(text_parts),
            "metadata": {
                "filename": filename,
                "rows": len(df),
                "columns": df.columns.tolist(),
                "dataframe": df,
            },
            "pages": 1,
            "dataframe": df,
        }

    def _process_docx(self, file_obj, filename: str) -> dict:
        """Extract text from Microsoft Word documents."""
        if not self._has_docx:
            raise ImportError("python-docx required for DOCX processing. Run: pip install python-docx")

        import docx

        if hasattr(file_obj, "read"):
            doc_bytes = io.BytesIO(file_obj.read())
        else:
            doc_bytes = io.BytesIO(file_obj)

        doc = docx.Document(doc_bytes)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            text_parts.append(f"\n--- Table {table_idx + 1} ---")
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    text_parts.append(row_text)

        return {
            "text": "\n".join(text_parts),
            "metadata": {
                "filename": filename,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            },
            "pages": 1,
        }
